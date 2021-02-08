import os
from xml.dom.minidom import getDOMImplementation
import re

import docx

from question import Question

DOCX_Q_DIR = "docx/questions/"
DOCX_A_DIR = "docx/answers/"
XML_DIR = "xml/"
MULTIPLE_CHOICE = ("A.", "B.", "C.", "D.", "E.")
TRUE_FALSE = ("1.", "2.")

for docx_file in os.listdir(DOCX_Q_DIR):
    q_content = []
    a_content = []
    question_list = []
    q_index = 1
    if docx_file != ".gitignore":
        file_path = DOCX_Q_DIR + docx_file
        name_arr = docx_file.split(" ")
        prefix = f"{name_arr[0].upper()}_{name_arr[1]}_Q_"
        for index, p in enumerate(docx.Document(file_path).paragraphs):
            if index == 0:
                pass
            else:
                if p.text != "":
                    if not (p.text.startswith(MULTIPLE_CHOICE) or p.text.startswith(TRUE_FALSE)):
                        q_name = f"{prefix}{q_index}"
                        q_content.append(q_name)
                        q_content.append(p.text)
                        q_index += 1
                    else:
                        opts = p.text
                        q_content.append(opts)

        for answer_file in os.listdir(DOCX_A_DIR):
            if f"{name_arr[0]} {name_arr[1]}" in answer_file:
                answer_file_path = DOCX_A_DIR + answer_file
                counter = 0
                for index, p in enumerate(docx.Document(answer_file_path).paragraphs):
                    if index == 0:
                        pass
                    else:
                        if p.text != "":
                            if counter <= 1:
                                a_content.append(p.text)
                                counter += 1
                            else:
                                counter = 0

        for i, e in enumerate(q_content):
            if not (e.startswith(MULTIPLE_CHOICE) or e.startswith(TRUE_FALSE) or e.startswith(
                    f"{name_arr[0].upper()}_{name_arr[1]}_Q_")):
                question_name = q_content[i - 1]
                question_text = e
                question_opts = {}
                question_type = ""
                if q_content[i + 1].startswith(MULTIPLE_CHOICE):
                    question_type = "multichoice"
                    question_opts["A"] = q_content[i + 1][3:]
                    question_opts["B"] = q_content[i + 2][3:]
                    question_opts["C"] = q_content[i + 3][3:]
                    question_opts["D"] = q_content[i + 4][3:]
                    question_opts["E"] = q_content[i + 5][3:]
                elif q_content[i + 1].startswith(TRUE_FALSE):
                    question_type = "truefalse"
                    question_opts["A"] = q_content[i + 1][3:]
                    question_opts["B"] = q_content[i + 2][3:]
                question = Question(question_name, question_text, question_type, question_opts)
                # checking for answers
                # need to isLower() and remove special chars
                temp_text = re.sub(r"[^a-zA-Z0-9]+", ' ', question_text.lower())
                if temp_text in [re.sub(r"[^a-zA-Z0-9]+", ' ', a.lower()) for a in a_content]:
                    a_index = [re.sub(r"[^a-zA-Z0-9]+", ' ', a.lower()) for a in a_content].index(temp_text)
                    question.set_ans(a_content[a_index + 1][3:])
                else:
                    # if answer still no found: take manual input
                    print("------------------")
                    print(f"404 not found:  {question_name} --- {question_text}")
                    answer = input(f"Enter answer manually: ")
                    print(f"you entered: {answer}")
                    question.set_ans(answer)
                    print("------------------")
                question_list.append(question)
                """
                ***ISSUE*** UNLIKE ALL OTHER COURSES, CNFAM XML FILES STILL HAVE `.DOC` IN THEIR NAME:
                    DOMAIN_03_QUESTION.doc.xml <-- ??
                """

        impl = getDOMImplementation()
        dom = impl.createDocument(None, "quiz", None)
        quiz_elem = dom.documentElement

        cat_question = dom.createElement("question")
        cat_question.setAttribute("type", "category")
        category = dom.createElement("category")
        text = dom.createElement("text")
        text.appendChild(dom.createTextNode(f"$course$/top/{name_arr[0]} {name_arr[1]}"))
        category.appendChild(text)
        cat_question.appendChild(category)
        quiz_elem.appendChild(cat_question)

        for question in question_list:
            question_elem = question.to_xml(dom)
            quiz_elem.appendChild(question_elem)

        with open(f"{XML_DIR}/{docx_file[0:23]}.xml", "w") as f:
            f.write(quiz_elem.toprettyxml())

